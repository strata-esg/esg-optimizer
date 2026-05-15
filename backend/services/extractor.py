"""
ESG Optimizer MVP - Extracteur de texte (PDF, DOCX, XLSX).
Extraction 100 % locale, pas de dépendance externe (remplace CloudConvert v1).
"""

import logging
from pathlib import Path

from backend.config import settings

logger = logging.getLogger(__name__)

# Formats acceptés
ALLOWED_EXTENSIONS = {"pdf", "docx", "xlsx"}


def _extract_pdf(file_path: str) -> str:
    """Extrait le texte d'un PDF via PyMuPDF (fitz)."""
    import fitz  # PyMuPDF

    text_parts: list[str] = []
    with fitz.open(file_path) as doc:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(f"--- Page {page_num} ---\n{page_text}")
    return "\n".join(text_parts)


def _extract_docx(file_path: str) -> str:
    """Extrait le texte d'un DOCX via python-docx (paragraphes + tableaux)."""
    from docx import Document

    doc = Document(file_path)
    text_parts: list[str] = []

    # Paragraphes
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Tableaux (souvent utilisés dans les rapports ESG)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def _extract_xlsx(file_path: str) -> str:
    """Extrait le texte d'un XLSX via openpyxl (toutes les feuilles)."""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    text_parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text_parts.append(f"=== Feuille : {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            row_values = [str(cell) for cell in row if cell is not None]
            if row_values:
                text_parts.append(" | ".join(row_values))

    wb.close()
    return "\n".join(text_parts)


# Dispatch par extension
_EXTRACTORS = {
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "xlsx": _extract_xlsx,
}


def extract_text(file_path: str, file_format: str) -> str:
    """
    Point d'entrée principal.
    Retourne le texte brut du document, tronqué à MAX_TEXT_LENGTH caractères.

    Raises:
        ValueError: si le format n'est pas supporté ou le fichier n'existe pas.
    """
    fmt = file_format.lower().strip(".")
    if fmt not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Format non supporté : '{fmt}'. Formats acceptés : {ALLOWED_EXTENSIONS}")

    path = Path(file_path)
    if not path.exists():
        raise ValueError(f"Fichier introuvable : {file_path}")

    logger.info("Extraction texte : %s (format=%s)", path.name, fmt)

    extractor = _EXTRACTORS[fmt]
    raw_text = extractor(file_path)

    # Troncature pour rester dans les limites du contexte GPT-4o
    max_len = settings.max_text_length
    if len(raw_text) > max_len:
        logger.warning(
            "Texte tronqué : %d -> %d caractères (%s)",
            len(raw_text), max_len, path.name,
        )
        raw_text = raw_text[:max_len]

    logger.info("Extraction terminée : %d caractères (%s)", len(raw_text), path.name)
    return raw_text
